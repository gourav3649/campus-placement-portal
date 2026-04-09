"""
Resume text extraction service.
Supports PDF and DOCX formats.
Includes validation and error handling.
"""
import os
import logging
from pathlib import Path
from typing import Optional
import pdfplumber
from docx import Document

logger = logging.getLogger(__name__)

# Minimum text length for a valid resume
MIN_RESUME_TEXT_LENGTH = 20


async def extract_resume_text(file_path: str) -> str:
    """
    Extract text from a resume file (PDF or DOCX).
    Validates extraction and returns non-empty text.
    
    Args:
        file_path: Full path to the resume file
        
    Returns:
        Extracted text content (guaranteed non-empty)
        
    Raises:
        FileNotFoundError: If file does not exist
        ValueError: If file format is unsupported or no text could be extracted
    """
    if not file_path or not isinstance(file_path, str):
        raise ValueError("File path must be a non-empty string")
    
    if not os.path.exists(file_path):
        logger.error(f"Resume file not found: {file_path}")
        raise FileNotFoundError(f"Resume file not found: {file_path}")
    
    file_ext = Path(file_path).suffix.lower()
    
    if file_ext == ".pdf":
        extracted_text = await _extract_pdf_text(file_path)
    elif file_ext == ".docx":
        extracted_text = await _extract_docx_text(file_path)
    else:
        raise ValueError(f"Unsupported file format: {file_ext}. Supported: .pdf, .docx")
    
    # Validate extracted text
    if not extracted_text or not extracted_text.strip():
        logger.warning(f"No text extracted from resume at {file_path}")
        raise ValueError("No text could be extracted from the resume file")
    
    # Ensure minimum text length
    if len(extracted_text.strip()) < MIN_RESUME_TEXT_LENGTH:
        logger.warning(
            f"Extracted resume text too short ({len(extracted_text.strip())} chars) "
            f"from {file_path}. Minimum: {MIN_RESUME_TEXT_LENGTH} chars"
        )
        raise ValueError(
            f"Extracted resume text is too short. "
            f"Minimum {MIN_RESUME_TEXT_LENGTH} characters required."
        )
    
    logger.debug(f"Resume text extracted successfully: {len(extracted_text)} chars")
    return extracted_text


async def _extract_pdf_text(file_path: str) -> str:
    """Extract text from PDF file using pdfplumber. Handles empty pages and errors gracefully."""
    try:
        text_parts = []
        page_count = 0
        empty_page_count = 0
        
        with pdfplumber.open(file_path) as pdf:
            page_count = len(pdf.pages)
            if page_count == 0:
                raise ValueError("PDF file is empty (no pages)")
            
            for page_num, page in enumerate(pdf.pages, 1):
                try:
                    text = page.extract_text()
                    if text and text.strip():
                        text_parts.append(text)
                    else:
                        empty_page_count += 1
                except Exception as page_error:
                    logger.warning(f"Failed to extract text from PDF page {page_num}: {str(page_error)}")
                    empty_page_count += 1
        
        if not text_parts:
            logger.error(f"PDF has {page_count} pages but no text could be extracted")
            raise ValueError(
                f"Could not extract text from any of the {page_count} pages in PDF. "
                f"The PDF may be image-based or corrupted."
            )
        
        extracted = "\n".join(text_parts).strip()
        logger.debug(
            f"PDF extraction complete: {page_count} pages, "
            f"{len(text_parts)} with text, {empty_page_count} empty. "
            f"Total text: {len(extracted)} chars"
        )
        return extracted
        
    except ValueError as ve:
        logger.error(f"PDF validation failed: {str(ve)}")
        raise
    except Exception as e:
        logger.error(f"Unexpected error extracting PDF: {str(e)}")
        raise Exception(f"Failed to extract text from PDF: {str(e)}")


async def _extract_docx_text(file_path: str) -> str:
    """Extract text from DOCX file using python-docx. Includes tables and handles empty documents."""
    try:
        doc = Document(file_path)
        text_parts = []
        para_count = len(doc.paragraphs)
        table_count = len(doc.tables)
        
        # Extract from paragraphs
        for para in doc.paragraphs:
            if para.text and para.text.strip():
                text_parts.append(para.text.strip())
        
        # Extract from tables
        table_rows = 0
        for table in doc.tables:
            table_rows += len(table.rows)
            for row in table.rows:
                for cell in row.cells:
                    if cell.text and cell.text.strip():
                        text_parts.append(cell.text.strip())
        
        if not text_parts:
            logger.error(
                f"DOCX has {para_count} paragraphs and {table_count} tables "
                f"({table_rows} rows) but no text could be extracted"
            )
            raise ValueError(
                f"Could not extract text from DOCX. "
                f"Document has {para_count} paragraphs and {table_count} tables but all are empty."
            )
        
        extracted = "\n".join(text_parts).strip()
        logger.debug(
            f"DOCX extraction complete: {para_count} paragraphs, "
            f"{len(text_parts)} text parts extracted, {table_count} tables with {table_rows} rows. "
            f"Total text: {len(extracted)} chars"
        )
        return extracted
        
    except ValueError as ve:
        logger.error(f"DOCX validation failed: {str(ve)}")
        raise
    except Exception as e:
        logger.error(f"Unexpected error extracting DOCX: {str(e)}")
        raise Exception(f"Failed to extract text from DOCX: {str(e)}")
